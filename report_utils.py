from docx import Document
from docx.shared import Pt
import io

def create_docx_from_markdown(text):
    """
    Converts simple Markdown text to a DOCX file object.
    Handling basic headers (#) and bullet points (-).
    """
    doc = Document()
    
    # Set default style if possible or just parse
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            # Simple ordered list check
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
            
    # Save to IO stream
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

from fpdf import FPDF

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Intelligence Report', 0, 1, 'C')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_pdf_from_markdown(text):
    """
    Converts simple Markdown text to a PDF file object.
    """
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Handle unicode characters broadly by assuming latin-1 for FPDF (limitation of basic FPDF)
    # properly we might need a utf-8 font, but for now let's try to encode/replace
    # or just clean the text
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(5)
            continue
            
        if line.startswith('# '):
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, line[2:].encode('latin-1', 'replace').decode('latin-1'), 0, 1)
            pdf.set_font("Arial", size=12)
        elif line.startswith('## '):
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, line[3:].encode('latin-1', 'replace').decode('latin-1'), 0, 1)
            pdf.set_font("Arial", size=12)
        elif line.startswith('### '):
            pdf.set_font("Arial", 'B', 13)
            pdf.cell(0, 10, line[4:].encode('latin-1', 'replace').decode('latin-1'), 0, 1)
            pdf.set_font("Arial", size=12)
        elif line.startswith('- ') or line.startswith('* '):
            pdf.cell(10) # Indent
            pdf.multi_cell(0, 5, chr(149) + " " + line[2:].encode('latin-1', 'replace').decode('latin-1'))
        elif line.startswith('1. '):
             pdf.cell(10)
             pdf.multi_cell(0, 5, line.encode('latin-1', 'replace').decode('latin-1'))
        else:
            pdf.multi_cell(0, 5, line.encode('latin-1', 'replace').decode('latin-1'))
            
    # Save to IO stream (as str for output logic, but streamlit needs bytes)
    # FPDF output() returns string in Py2, bytes/string in Py3 depending on args.
    # dest='S' returns as string.
    
    pdf_out = pdf.output(dest='S').encode('latin-1')
    return io.BytesIO(pdf_out)
